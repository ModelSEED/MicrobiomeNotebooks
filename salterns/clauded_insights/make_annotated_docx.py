#!/usr/bin/env python
"""
Write a copy of the manuscript DOCX with visible in-line annotations (red, bold,
"CLAUDE >>") placed immediately after each passage that a change affects.  The
original figures/formatting are preserved (python-docx round-trip).  Annotations
reference the corrected figures and numbers in this folder; nothing in the author's
own text is altered -- the annotations are added paragraphs only.
"""
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

SRC = Path(__file__).resolve().parents[1] / "Metabolic Manuscript.docx"
OUT = Path(__file__).resolve().parent / "Metabolic Manuscript (clauded annotations).docx"
RED = RGBColor(0xC0, 0x00, 0x00)


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    run = np.add_run("CLAUDE »  " + text)
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(9)
    return np


# (substring to find in a paragraph, annotation text)
ANNOTATIONS = [
 ("normalized by DNA density for all modeling efforts",
  "DNA normalization is applied to BOTH the methane axis and (elsewhere) the abundance axis. "
  "Dividing two quantities by a shared denominator (DNA) inflates their correlation "
  "(spurious correlation of ratios). Report the raw-areal-unit result as primary and state the "
  "normalization explicitly. See numbers_to_reconcile.md and figure7_corrected.png."),
 ("correlated remarkably well (R2 = 0.966)",
  "STALE NUMBER. R2 = 0.966 comes from a pre-DNA-weighting analysis and matches nothing in the "
  "current data. Submitted Figure 7 shows R2 = 0.89 (DNA-normalized); on a consistent raw-areal "
  "basis it is R2 = 0.37. Honest statistic: Spearman rho = 0.60, p = 0.21 (n = 6, not "
  "significant). Corrected values in figure7_corrected.png / figure7_stats.json."),
 ("almost identical fit to the best fitting regression (R2 = 0.968)",
  "Lead with the UNCONSTRAINED (parameter-free) fit so the ratio cannot be read as tuned to the "
  "answer. Genus breakdown of the modelled MePn methane: Roseovarius 77%, Albimonas 11% "
  "(family_contribution.png) -- good support for this sentence."),
 ("correlated slightly less well (R2 = 0.862)",
  "STALE. On the corrected raw basis the betaine model has NO relationship with measured methane "
  "(Spearman rho = -0.09, R2 = 0.10). 'slightly less well' overstates it -- it is essentially "
  "uncorrelated. See figure7_corrected.png panel a."),
 ("R2 = 0.974 for all samples and R2 = 0.966 for unrestored samples",
  "STALE / KEY CORRECTION. Recomputed on raw units over all 12 cores: MePn consumers rho = 0.73 "
  "(p = 0.007) and betaine consumers rho = 0.70 (p = 0.011) -- BOTH significant and of similar "
  "strength. The 'MePn >> betaine' separation is an artifact of the n = 6 subset + DNA "
  "double-normalization. On the raw unrestored-only subset neither is significant. See "
  "figure7_corrected.png panel b and numbers_to_reconcile.md."),
 ("reproduced the experimental methane emissions with high accuracy (R2 = 0.978)",
  "STALE (R2 = 0.978 unsupported) AND internal error: the preceding 'R2 = 0.798' is the OLD "
  "betaine unrestored-abundance number, mis-attached to MePn organisms. Corrected framing: the "
  "model reproduces the RANK ORDER (rho = 0.60, n = 6) and is best treated as a consistency "
  "check alongside the independent culturing (Fig 6) and substrate-addition (Fig 4) evidence."),
 ("Assessing methylphosphonate and betaine metabolisms as potentially explaining",
  "REPLACE Figure 7 with figure7_corrected.png/.svg: linear fit + log axes retained; trendlines "
  "rendered correctly (were broken piecewise curves); a genuine 1:1 line added; all 12 cores "
  "shown coloured by restoration; Spearman rho + p + n reported. Add Figure 7c = "
  "family_contribution.png (genus breakdown of modelled methane)."),
 ("Co-occurrence networks that explore the relationships",
  "REPLACE Figure 8 with figure8_relabeled.png/.svg: same Graphical-Lasso graph and Louvain "
  "communities (Q = 0.57), but labels de-collided (the submitted panel overlaps labels such as "
  "'BM003KS3-K002'). Add the caveat that the precision matrix is inferred from n = 12 over ~140 "
  "taxa (p >> n); edges are exploratory -- bootstrap edge-stability recommended."),
 ("For the one methanogen MAG, a methanogenic",
  "Note for reviewers: the dereplicated MAG set used for abundance-weighting is entirely "
  "bacterial (280 Bacteria, 0 archaea); this one methanogen MAG is not in that table, so the "
  "betaine->TMA->methane route is closed with an empirical 0.06 x TMA scalar (from Fig 4) rather "
  "than a modelled archaeon. State this explicitly so the archaeal step is not assumed modelled."),
]

doc = Document(str(SRC))
# banner at the very top
first = doc.paragraphs[0]
banner = insert_after(first,
    "Annotated copy (auto-generated). Red 'CLAUDE >>' notes flag changes proposed after review "
    "of the code, data, and figures; see clauded_insights/README.md, numbers_to_reconcile.md, "
    "figure7_corrected.png, family_contribution.png, and figure8_relabeled.png. The author's "
    "original text is unchanged; only annotation paragraphs were added.")

hits, misses = 0, []
for needle, note in ANNOTATIONS:
    for p in doc.paragraphs:
        if needle in p.text:
            insert_after(p, note)
            hits += 1
            break
    else:
        misses.append(needle[:50])

doc.save(str(OUT))
print(f"annotations placed: {hits}/{len(ANNOTATIONS)}")
if misses:
    print("NOT FOUND:", misses)
print("saved", OUT.name)
